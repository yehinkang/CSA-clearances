'''
Created on Nov 15, 2018
@author: jdong

This module is intended for documentation of the TLBU Tools. It provides a few interfaces for creating a WORD
dpcument by warpping the python-docx APIs.

MODIFICATIONS:
5/3/2022: _add_table def: if there are less cells in a row, the last cell will merge with rest of the cells in same row

'''
import os
from datetime import datetime
import warnings

import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.api import Document as _Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Length
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches

# define the default_docx package for using pyInstaller to create a onefile bundle
default_docx = os.path.join(os.getcwd(), 'apitlbu', 'resources', 'templates_docx', 'default.docx')


class ReportDocx:
    file_docx = None
    report_title = 'Title of the Report'
    img_cover = None

    def __init__(self, file_path, report_title=None, img_cover=None, append=False):
        '''
        Parameters:
            file_path: a full path for the WORD to be saved as.
            report_tile: title of the entire document
            image_cover: an image to be included in the report cover. Date and Time info will be automatically
            added below the image.
            document: if not None, will be opened for modification
        '''

        # create a document object; add needed styles, e.g. font family and size
        self.file_docx = file_path

        if append and os.path.isfile(file_path):
            self.document = _Document(file_path)
        else:
            self.document = _Document(default_docx)
        if report_title:
            self.report_title = report_title
        if img_cover:
            self.img_cover = img_cover

        styles = self.document.styles

        style_norm_cour = styles.add_style('norm_cour', WD_STYLE_TYPE.PARAGRAPH)
        font = style_norm_cour.font
        font.name = 'Courier New'
        font.size = Pt(10)
        #font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        font.color.rgb = RGBColor(0x00, 0x00, 0x00) # black
        self.style_norm = style_norm_cour
        
        style_norm_cour_red = styles.add_style('norm_cour_red', WD_STYLE_TYPE.PARAGRAPH)
        font_red = style_norm_cour_red.font 
        font_red.name = 'Courier New'
        font_red.size = Pt(10)
        font_red.color.rgb = RGBColor(0xFF, 0x00, 0x00) # red
        self.style_norm_red = style_norm_cour_red
        #document.save(self.file_docx)

    def create_report_cover(self, url=None, img_width=4.5, app_name=None):
        # if file exists, overwire?
        '''
        Parameter url: if defined, a link will be added with display text "TLBU Home".
        Call this method to create a cover page using the title and image info defined for the class,
        and then add the cover page to the document as the first page.

        url: link to the text 'ENGINEERING DESIGN SUITE'
        img_width: width of the image in inches
        app_name: name of the application
        '''

        self.document.add_heading(self.report_title, 0)

        if self.img_cover:
            self.document.add_picture(self.img_cover, width=Inches(img_width))
            last_paragraph = self.document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = self.document.add_paragraph()
        p.style = self.style_norm

        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        # time information

        p.add_run().add_break(WD_BREAK.LINE)
        if app_name:
            p.add_run(f'Application: ').bold = True
            p.add_run(f'{app_name}').bold = False
            p.add_run().add_break(WD_BREAK.LINE)

        p.add_run('\nReport is saved as: ').bold = True
        p.add_run(f'{os.path.basename(self.file_docx)}').bold = False
        p.add_run().add_break(WD_BREAK.LINE)

        string_time = datetime.utcnow().strftime('%H:%M:%S %m/%d/%Y')
        p.add_run('\n\tAt UTC time ' + string_time + '\n').bold = False
        p.add_run().add_break(WD_BREAK.LINE)

        p.add_run('Powered by T.L.B.U. ').bold = True
        if url:
            _add_hyperlink(p, 'ENGINEERING DESIGN SUITE', url=url)
        else:
            p.add_run('ENGINEERING DESIGN SUITE').bold = True

        current_section = self.document.sections[-1]

        # turn on "different first page" to surpress header on cover page
        current_section.different_first_page_header_footer = True
        footer_0 = current_section.first_page_footer


        self.document.save(self.file_docx)

    def insert_page_break(self):
        '''
        Call this method to create a hard page break
        '''
        p = self.document.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        self.document.save(self.file_docx)

    def write_report_section(self, report_data, section_title, orient=0, dimension=0,
                             widths=None,
                             space_pt_before=0, space_pt_after=0,
                             footer_left=None, footer_center=None, footer_right=None):
        '''
        Call this method to create a section with content and add it to the document.

        report_data: a list of dictionaries. Each dictionary is a sub-section and contains the following:
            is_page_break: boolean.
                If True, will add a page break
                If None or False, will be ignored;
            istable: boolean. Content will be treated and displayed as a table if it is set to True

            title - title of this sub-section

            header:  applicable if is_table is True. A list of text representing column header of a table;

            font_size: optional, integer representing font size in points

            content:
                if istable is True: a two-dimensional list. each item represents one table cell

                if is_table is False:
                    a list of dictionaries representing paragraphs to be written into this section. These following
                    keys can be specified in each list item (a dictionary):
                        bold: boolean. Bold fonts will be used if set to True
                        indent_l: integer, representing paragraph indent in pts
                        text: content of the paragraph
                        text-align: a text paragraph with alignment specified. Possible aligntments can be one of these:
                            center, left, right
                        color: red or None

            tablefooter: applicable if is_table is True. A list of text represent a list of footers of the table

            col_width: applicable is is_table is True. A list of float numbers representing column width in Inches for
                       specific table [widths associated with section is not recommended]

        section_title: a string that represents the title of this section
        orient: optional integer: 0-portrait; 1-landscape (applied to this sectioin only)
        dimension: optional integer: 0 - LETTER; 1-TOBLOID (applied to this section only)
        widths: applicable if is_table is True. A list of integers representing column width in Inches
                [Depracted; replaced with col_width item in each content for individual table. See "content"]
        space_pt_before: optional integer, space in point before each paragraph
        space_pt_after: optional integer, space in point after each paragraph
        footer_left: optional text - left portion of footer for this section
        footer_center: optional text - center portion of footer for this section
        footer_right: optional text - right portion of footer for this section
        '''

        '''
        Add deprecation warning...
        '''

        if widths:
            warnings.warn('Key "widths" for section is deprecated. Use "col_width" for each individual table instead.')
        #self.document.add_page_break()

        # check parameter validity
        if orient > 1:
            raise ValueError('Value for Orientation has to be one of the two values [0, 1]')
        if dimension > 1:
            raise ValueError('Value for Dimension has to be one of the two values [0, 1]')

        self.document.add_section()
        current_section = self.document.sections[-1]

        _add_section_footer(document=self.document, section=current_section,
                            footer_left=footer_left, footer_center=footer_center, footer_right=footer_right,
                            orientation=orient, dimension=dimension)

        # ------------------------------
        # Page dimensions and orientation:
        # section.orientation, section.page_width, section.page_height
        # (PORTRAIT(0), 7772400, 10058400)  # (Inches(8.5), Inches(11))
        # ------------------------------
        if dimension == 1:
            current_section.page_height = Inches(17)
            current_section.page_width = Inches(11)

        if orient == 1:
            height = current_section.page_height
            width = current_section.page_width
            new_width, new_height = max(height, width), min(height, width)
            current_section.page_width = new_width
            current_section.page_height = new_height
            current_section.orientation = WD_ORIENT.LANDSCAPE

        self.document.add_heading(section_title, level=1)

        styles_in_doc = self.document.styles
        styles_paragraph = [s.name for s in styles_in_doc if s.type == WD_STYLE_TYPE.PARAGRAPH]

        for i, data in enumerate(report_data):
            if data:
                if data.get('is_page_break'):
                    self.insert_page_break()

                if data.get('istable'):
                    _add_table(self.document, data, self.style_norm)
                    if 'tablefooter' in data:
                        _add_table_footer(self.document, data['tablefooter'], self.style_norm)
                        
                else:
                    font_size = data.get('font_size')
                    if font_size:
                        font_normal = self.style_norm.font
                        font_normal.size = Pt(font_size)

                        font_red = self.style_norm_red.font
                        font_red.size = Pt(font_size)

                    # add paragraph
                    _add_paragraph(self.document, data,
                                   style_normal=self.style_norm, style_cour_red=self.style_norm_red,
                                   space_pt_before=space_pt_before, space_pt_after=space_pt_after)

        self.document.save(self.file_docx)

    def insert_image(self, img_path, inch_width=1.5, para_comment=None, orient=0):
        if img_path is None:
            return
        self.document.add_picture(img_path, width=Inches(inch_width))
        if (para_comment):
            _add_paragraph(self.document, para_comment, self.style_norm, style_cour_red=self.style_norm_red)
        self.document.save(self.file_docx)

    def write_report_section_old(self, report_txt, section_title, orient=0):
        '''
        Deprecated
        '''
        self.document.add_page_break()
        self.document.add_section()
        current_section = self.document.sections[-1]
        if orient == 1:
            new_width, new_height = current_section.page_height, current_section.page_width
            current_section.page_width = new_width
            current_section.page_height = new_height
            current_section.orientation = WD_ORIENT.LANDSCAPE
        
        self.document.add_heading(section_title, level=1)
 
        for text in report_txt:
            p = self.document.add_paragraph()
            p.style = self.style_norm
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text['content']).bold = text['bold']
            #font = run.font
            #font.bold = text['bold']
        self.document.save(self.file_docx)


def _add_section_footer(document, section, footer_left, footer_center, footer_right,
                        orientation=0, dimension=0):
    # tunn off "different first page" to include footer for all pages in this section
    section.different_first_page_header_footer = False

    if footer_left is None and footer_center is None and footer_right is None:
        # No footer is defined for this section
        return

    # reserver 117 chars (consolas size 10) for footer: portrait tabloid
    # reserver 188 chars (consolas size 10) for footer: landscape tabloid
    # reserver 85 chars (consolas size 10) for footer: portrait letter
    # reserver 111 chars (consolas size 10) for footer: landscape letter

    # left: 45; center:100; right: 43
    text_left = ' ' if footer_left is None else footer_left
    text_center = ' ' if footer_center is None else footer_center
    text_right = ' ' if footer_right is None else footer_right
    #text_left = '{:<45}'.format(text_left)
    #text_center = '{:^100}'.format(text_center)
    #text_right = '{:>43}'.format(text_right)

    if orientation == 0 and dimension == 0:
        # left: length of string=20; truncated to 15 chars
        text_left = '{:<20.15}'.format(text_left)
        text_center = '{:^38.30}'.format(text_center)
        text_right = '{:>20.15}'.format(text_right)
        text_footer = f'{text_left}{text_center}{text_right}'
    elif orientation == 1 and dimension == 0:
        # Letter size, landscape
        text_left = '{:<30.25}'.format(text_left)
        text_center = '{:^50.40}'.format(text_center)
        text_right = '{:>30.25}'.format(text_right)
        text_footer = f'{text_left}{text_center}{text_right}'
    elif orientation == 0 and dimension == 1:
        # tabloid size, portrait
        text_left = '{:<33.25}'.format(text_left)
        text_center = '{:^50.40}'.format(text_center)
        text_right = '{:>33.25}'.format(text_right)
        text_footer = f'{text_left}{text_center}{text_right}'
    elif orientation == 1 and dimension == 1:
        # Tabloid size, landscape
        text_left = '{:<45.25}'.format(text_left)
        text_center = '{:^100.40}'.format(text_center)
        text_right = '{:>43.25}'.format(text_right)
        text_footer = f'{text_left}{text_center}{text_right}'
    else:
        text_footer = f'{text_left}\t{text_center}\t{text_right}'

    list_style_names = [style.name for style in document.styles]
    if 'consolas' in list_style_names:
        style_consolas = document.styles['consolas']

    else:
        style_consolas = document.styles.add_style('consolas', WD_STYLE_TYPE.PARAGRAPH)

        font = style_consolas.font
        font.name = 'Consolas'
        font.size = Pt(10)

    #print(section.different_first_page_header_footer)

    footer_section = section.footer
    footer_paragraph = footer_section.paragraphs[0]

    footer_paragraph.text = text_footer
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.style = style_consolas
    #footer_paragraph.style = document.styles['Footer']


def _add_table_footer(document, list_footer, p_style):
    p = document.add_paragraph()
    p.style = p_style
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('Note(s):').bold = False

    for footer in list_footer:
        #p = document.add_paragraph(footer, style='List Number')
        p = document.add_paragraph()
        p.style = 'List Number'

        p.style.font.name = 'Courier New'
        run = p.add_run(footer).bold = False
        
        
def _add_table(document, data_table, style_normal):

    title = data_table['title']
    header = data_table['header']
    content = data_table['content']
    col_width = data_table.get('col_width')

    # add title
    p = document.add_paragraph()
    p.style = style_normal
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title).bold = True
                
    ### add a table ###
    n_cols = len(header)

    table = document.add_table(rows=1, cols=n_cols)

    # populate header row
    heading_cells = table.rows[0].cells
    for i, h in enumerate(header):
        
        heading_cells[i].text = h
        p1 = heading_cells[i].paragraphs[0]
        p1.style = style_normal
        #p1.add_run().bold = True
                
    # populate cell data_table
    for row_data in content:
        cells = table.add_row().cells

        """
        if len(cells) < len(row_data):
            # cells created per row is shorter than row_data::: something wroing
            print('>>> when writing to docx, row_data is a longer list than cells-for-row:')
            print('>>> length of cells=', len(cells), 'row_data=', row_data)
            continue"""

        for i, cell in enumerate(row_data):
            if i < n_cols:
                cells[i].text = cell
                p2 = cells[i].paragraphs[0]
                p2.style = style_normal
            else:
                cells[n_cols-1].text = 'DATAERROR'
                p2 = cells[n_cols-1].paragraphs[0]
                p2.style = style_normal
                continue
            #p2.add_run().bold = False

        # merge with rest of cells in the row if less cells specified
        if i < n_cols - 1:
            cell = table.rows[-1].cells[i]
            for j in range(i+1, n_cols):
                cell.merge(table.rows[-1].cells[j])

    table.style = 'Medium Shading 1 Accent 1'

    if col_width:
        print(f'::: use fixed width for table: {col_width}')
        table.allow_autofit = False
        inch_widths = []
        for width in col_width:
            inch_widths.append(Inches(width))

        for row in table.rows:
            for idx, width in enumerate(inch_widths):
                if idx < n_cols:
                    row.cells[idx].width = width
    else:
        table.allow_autofit = True
    
    
def _add_paragraph(document, data_para, style_normal, space_pt_before=0, space_pt_after=0, style_cour_red=None):
    title = data_para.get('title')
    content = data_para.get('content')

    # add title
    p = document.add_paragraph()
    p.style = style_normal

    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_pt_before)
    p.paragraph_format.space_after = Pt(space_pt_after)

    p.add_run().add_break(WD_BREAK.LINE)
    if title:
        run = p.add_run(title).bold = True

    # add content, a list of paragraph dictionaries, and color/font/etc attributes
    # dictionary keys: 'text', 'color', 'bold', 'italic', 'underl'
    for para in content:
        if 'text_align' in para:
            if para['text_align'].lower() == 'left':
                text_alignment = WD_ALIGN_PARAGRAPH.LEFT

            elif para['text_align'].lower() == 'center':
                text_alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                text_alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            text_alignment = WD_ALIGN_PARAGRAPH.LEFT

        try:
            indent_left = int(para['indent_l']) if 'indent_l' in para else 0
            indent_right = int(para['indent_r']) if 'indent_r' in para else 0
            is_bold = int(para['bold']) if 'bold' in para else False
        except:
            print('>>> key error')

        p1 = document.add_paragraph()
        
        if para.get('color') is None:
            p1.style = style_normal
        elif para.get('color') == 'red':
            p1.style = style_cour_red
        else:
            p1.style = style_normal

        p1.paragraph_format.alignment = text_alignment
        p1.paragraph_format.left_indent = Pt(indent_left)
        p1.paragraph_format.right_indent = Pt(indent_right)
        p1.paragraph_format.space_before = Pt(space_pt_before)
        p1.paragraph_format.space_after = Pt(space_pt_after)

        run = p1.add_run(para['text']).bold = is_bold


def _get_cover_text():
    text = []
    text.append('Due to the large quantity of SFD\'s that are to ' \
        'be installed within each span, it is not possible to apply ' \
        'them to the PLS-CADD model as concentrated loads.  As a '\
        'result, hand calculations have been performed (see Appendix G) ' \
        'to determine an \"effective K-factor\" for the shieldwire and OPGW '\
        'to account for the additional weight on each conductor due to ice '\
        'accumulation on the SFD\'s.  These effective K-factors are then used '\
        'to adjust the wire tension for each loading condition associated with '\
        'ice on the wire.  Hand calculations were also performed to determine a '\
        'wind pressure adjustment factor to account for wind on the iced SFD\'s.  '\
        'Manufacturer\'s literature suggests that ice can develop on only 50% of '\
        'the surface area of the SFD.') 
    text.append('Wind and weight effects due to the application of SFD\'s '\
        'for non-ice load cases are considered negligible')
    
    return text


def _add_hyperlink(paragraph, text, url):
    #https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

